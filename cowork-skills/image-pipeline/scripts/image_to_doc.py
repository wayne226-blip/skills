#!/usr/bin/env python3
"""Bridge between AI-generated images and document pipelines.

Takes image files (from Nano Banana, DALL-E, local files, or URLs)
and inserts them into PPTX, DOCX, or PDF (via HTML for WeasyPrint).

Usage:
    # Insert image into PPTX slide
    python image_to_doc.py image.png output.pptx --slide 2 --position center

    # Insert image into DOCX
    python image_to_doc.py image.png output.docx --width 5

    # Generate HTML img tag for WeasyPrint PDF
    python image_to_doc.py image.png --html --width 400px

    # Download from URL first, then insert
    python image_to_doc.py --url "https://example.com/photo.jpg" output.pptx

    # Batch: insert multiple images
    python image_to_doc.py img1.png img2.png img3.png output.pptx --layout grid
"""

import argparse
import base64
import os
import sys
import shutil
import subprocess
import urllib.request
from pathlib import Path


def image_to_base64(image_path):
    """Convert an image file to a base64 data URI."""
    ext = Path(image_path).suffix.lower()
    mime_map = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".svg": "image/svg+xml",
        ".webp": "image/webp",
    }
    mime = mime_map.get(ext, "image/png")

    with open(image_path, "rb") as f:
        data = base64.b64encode(f.read()).decode("utf-8")

    return f"{mime};base64,{data}"


def download_image(url, dest_dir="/tmp"):
    """Download an image from a URL and return the local path."""
    filename = url.split("/")[-1].split("?")[0] or "downloaded_image.png"
    dest = os.path.join(dest_dir, filename)
    urllib.request.urlretrieve(url, dest)
    print(f"Downloaded: {dest}")
    return dest


def get_image_dimensions(image_path):
    """Get image dimensions using Python (Pillow)."""
    try:
        from PIL import Image
        with Image.open(image_path) as img:
            return img.size  # (width, height)
    except ImportError:
        # Fallback: try identify command
        try:
            result = subprocess.run(
                ["identify", "-format", "%w %h", image_path],
                capture_output=True, text=True
            )
            if result.returncode == 0:
                w, h = result.stdout.strip().split()
                return (int(w), int(h))
        except FileNotFoundError:
            pass
    return (800, 600)  # Sensible default


def insert_into_pptx(image_paths, output_path, slide_num=None, position="center",
                      layout="single", width_inches=None, template=None):
    """Insert images into a PowerPoint file using PptxGenJS."""
    images_data = []
    for img in image_paths:
        b64 = image_to_base64(img)
        w, h = get_image_dimensions(img)
        images_data.append({"path": img, "base64": b64, "width": w, "height": h})

    # Calculate sizing
    slide_w, slide_h = 10, 5.625  # 16:9

    if layout == "grid" and len(images_data) > 1:
        cols = 2 if len(images_data) <= 4 else 3
        rows = (len(images_data) + cols - 1) // cols
        cell_w = (slide_w - 1.2) / cols
        cell_h = (slide_h - 1.5) / rows

        placements = []
        for i, img_data in enumerate(images_data):
            r, c = divmod(i, cols)
            x = 0.6 + c * cell_w
            y = 0.8 + r * cell_h
            aspect = img_data["width"] / img_data["height"]
            if aspect > cell_w / cell_h:
                w = cell_w * 0.9
                h = w / aspect
            else:
                h = cell_h * 0.9
                w = h * aspect
            placements.append({"x": x, "y": y, "w": w, "h": h, "data": img_data["base64"]})
    else:
        placements = []
        for img_data in images_data:
            aspect = img_data["width"] / img_data["height"]
            if width_inches:
                w = width_inches
                h = w / aspect
            else:
                max_w, max_h = slide_w - 1.2, slide_h - 1.5
                if aspect > max_w / max_h:
                    w = max_w
                    h = w / aspect
                else:
                    h = max_h
                    w = h * aspect

            if position == "center":
                x = (slide_w - w) / 2
                y = (slide_h - h) / 2
            elif position == "left":
                x = 0.6
                y = (slide_h - h) / 2
            elif position == "right":
                x = slide_w - w - 0.6
                y = (slide_h - h) / 2
            else:
                x = (slide_w - w) / 2
                y = (slide_h - h) / 2

            placements.append({"x": x, "y": y, "w": w, "h": h, "data": img_data["base64"]})

    # Generate JS
    js_lines = [
        'const pptxgen = require("pptxgenjs");',
        'const pres = new pptxgen();',
        'pres.layout = "LAYOUT_16x9";',
    ]

    for i, pl in enumerate(placements):
        js_lines.append(f'const slide{i} = pres.addSlide();')
        js_lines.append(
            f'slide{i}.addImage({{ data: "{pl["data"]}", '
            f'x: {pl["x"]:.2f}, y: {pl["y"]:.2f}, w: {pl["w"]:.2f}, h: {pl["h"]:.2f} }});'
        )

    js_lines.append(f'pres.writeFile({{ fileName: "{output_path}" }}).then(() => {{')
    js_lines.append(f'  console.log("PPTX created: {output_path}");')
    js_lines.append('});')

    js_path = "/tmp/_img_to_pptx.js"
    with open(js_path, "w") as f:
        f.write("\n".join(js_lines))

    result = subprocess.run(["node", js_path], capture_output=True, text=True)
    if result.returncode != 0:
        print(f"Error: {result.stderr}", file=sys.stderr)
        sys.exit(1)
    print(result.stdout.strip())


def insert_into_docx(image_paths, output_path, width_inches=5.0):
    """Insert images into a Word document."""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    for img_path in image_paths:
        doc.add_picture(img_path, width=Inches(width_inches))
        doc.add_paragraph()  # spacing

    doc.save(output_path)
    print(f"DOCX created: {output_path} ({os.path.getsize(output_path):,} bytes)")


def generate_html_tag(image_paths, width="100%", css_class=""):
    """Generate HTML img tags for WeasyPrint PDF pipeline."""
    tags = []
    for img_path in image_paths:
        b64 = image_to_base64(img_path)
        cls = f' class="{css_class}"' if css_class else ''
        tags.append(
            f'<img src="data:{b64}"{cls} '
            f'style="width: {width}; display: block; margin: 1em auto;" '
            f'alt="{Path(img_path).stem}">'
        )
    return "\n".join(tags)


def main():
    parser = argparse.ArgumentParser(description="Insert images into documents")
    parser.add_argument("images", nargs="*", help="Image file paths")
    parser.add_argument("output", nargs="?", help="Output file (.pptx, .docx)")
    parser.add_argument("--url", help="Download image from URL first")
    parser.add_argument("--slide", type=int, help="Target slide number (PPTX)")
    parser.add_argument("--position", default="center", choices=["center", "left", "right"])
    parser.add_argument("--layout", default="single", choices=["single", "grid"])
    parser.add_argument("--width", type=float, help="Image width in inches")
    parser.add_argument("--html", action="store_true", help="Output HTML img tags instead")
    parser.add_argument("--html-width", default="100%", help="CSS width for HTML output")
    parser.add_argument("--template", help="PPTX reference template")
    args = parser.parse_args()

    # Gather images
    images = list(args.images or [])
    if args.url:
        images.append(download_image(args.url))

    if not images:
        parser.error("No images specified")

    # Verify images exist
    for img in images:
        if not os.path.exists(img):
            print(f"Error: Image not found: {img}", file=sys.stderr)
            sys.exit(1)

    if args.html:
        print(generate_html_tag(images, width=args.html_width))
        return

    if not args.output:
        parser.error("Output file required (unless using --html)")

    ext = Path(args.output).suffix.lower()
    if ext == ".pptx":
        insert_into_pptx(images, args.output, slide_num=args.slide,
                         position=args.position, layout=args.layout,
                         width_inches=args.width, template=args.template)
    elif ext == ".docx":
        insert_into_docx(images, args.output, width_inches=args.width or 5.0)
    else:
        print(f"Unsupported output format: {ext}", file=sys.stderr)
        print("Use .pptx, .docx, or --html for PDF pipeline", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
