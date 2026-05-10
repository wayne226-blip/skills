#!/usr/bin/env python3
"""Generate a branded DOCX reference template for Pandoc.

Creates a .docx with Wayne's brand fonts, colours, heading styles,
and page setup. Use with Pandoc's --reference-doc flag to get
consistent branded Word output.

Usage:
    python create_docx_template.py [output.docx] [--palette NAME]

Palettes: midnight (default), forest, coral, ocean, charcoal
"""

import argparse
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

PALETTES = {
    "midnight": {
        "primary": RGBColor(0x1E, 0x27, 0x61),
        "secondary": RGBColor(0x2B, 0x6C, 0xB0),
        "text": RGBColor(0x2D, 0x37, 0x48),
        "muted": RGBColor(0x71, 0x80, 0x96),
        "heading_font": "Georgia",
        "body_font": "Calibri",
    },
    "forest": {
        "primary": RGBColor(0x2C, 0x5F, 0x2D),
        "secondary": RGBColor(0x97, 0xBC, 0x62),
        "text": RGBColor(0x1A, 0x20, 0x2C),
        "muted": RGBColor(0x6B, 0x72, 0x80),
        "heading_font": "Georgia",
        "body_font": "Calibri",
    },
    "coral": {
        "primary": RGBColor(0x2F, 0x3C, 0x7E),
        "secondary": RGBColor(0xF9, 0x61, 0x67),
        "text": RGBColor(0x2D, 0x37, 0x48),
        "muted": RGBColor(0x71, 0x80, 0x96),
        "heading_font": "Georgia",
        "body_font": "Calibri",
    },
    "ocean": {
        "primary": RGBColor(0x06, 0x5A, 0x82),
        "secondary": RGBColor(0x1C, 0x72, 0x93),
        "text": RGBColor(0x1A, 0x20, 0x2C),
        "muted": RGBColor(0x64, 0x74, 0x8B),
        "heading_font": "Georgia",
        "body_font": "Calibri",
    },
    "charcoal": {
        "primary": RGBColor(0x36, 0x45, 0x4F),
        "secondary": RGBColor(0x21, 0x21, 0x21),
        "text": RGBColor(0x21, 0x21, 0x21),
        "muted": RGBColor(0x6B, 0x72, 0x80),
        "heading_font": "Arial",
        "body_font": "Arial",
    },
}


def create_template(output_path, palette_name):
    p = PALETTES.get(palette_name, PALETTES["midnight"])
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    # Style: Normal (body text)
    style = doc.styles["Normal"]
    style.font.name = p["body_font"]
    style.font.size = Pt(11)
    style.font.color.rgb = p["text"]
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Style: Title
    title_style = doc.styles["Title"]
    title_style.font.name = p["heading_font"]
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = p["primary"]
    title_style.font.bold = True
    title_style.paragraph_format.space_after = Pt(4)

    # Style: Subtitle
    subtitle_style = doc.styles["Subtitle"]
    subtitle_style.font.name = p["body_font"]
    subtitle_style.font.size = Pt(14)
    subtitle_style.font.color.rgb = p["secondary"]
    subtitle_style.font.italic = True
    subtitle_style.paragraph_format.space_after = Pt(12)

    # Heading styles
    for level, (size, bold) in enumerate(
        [(22, True), (16, True), (13, True), (11, True)], start=1
    ):
        style_name = f"Heading {level}"
        h_style = doc.styles[style_name]
        h_style.font.name = p["heading_font"]
        h_style.font.size = Pt(size)
        h_style.font.bold = bold
        h_style.font.color.rgb = p["primary"] if level <= 2 else p["secondary"]
        h_style.paragraph_format.space_before = Pt(18 if level == 1 else 14)
        h_style.paragraph_format.space_after = Pt(6)

    # Add sample content so Pandoc can read the styles
    doc.add_paragraph("Document Title", style="Title")
    doc.add_paragraph("Subtitle text", style="Subtitle")
    doc.add_heading("Heading 1", level=1)
    doc.add_paragraph("Body text with Normal style.")
    doc.add_heading("Heading 2", level=2)
    doc.add_paragraph("More body text.")
    doc.add_heading("Heading 3", level=3)
    doc.add_paragraph("Sub-section content.")

    doc.save(output_path)
    file_size = os.path.getsize(output_path)
    print(f"Template created: {output_path} ({file_size:,} bytes, palette: {palette_name})")
    print(f"Styles: Title, Subtitle, Heading 1-4, Normal")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate branded DOCX reference template")
    parser.add_argument("output", nargs="?", default="wayne_brand_template.docx", help="Output file")
    parser.add_argument("--palette", default="midnight", choices=list(PALETTES.keys()))
    args = parser.parse_args()
    create_template(args.output, args.palette)
