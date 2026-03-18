"""
cv_to_docx.py — bundled docx generator for cv-builder skill
Usage: python3 cv_to_docx.py input.md output.docx
"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1a, 0x1a, 0x2e)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(0)
    return p


def parse_md(md_text):
    """Parse markdown into a list of (type, content) tuples."""
    lines = md_text.splitlines()
    blocks = []
    for line in lines:
        if line.startswith('# '):
            blocks.append(('h1', line[2:].strip()))
        elif line.startswith('## '):
            blocks.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            blocks.append(('h3', line[4:].strip()))
        elif line.startswith('- ') or line.startswith('* '):
            blocks.append(('bullet', line[2:].strip()))
        elif line.startswith('---'):
            blocks.append(('hr', ''))
        elif line.strip() == '':
            blocks.append(('blank', ''))
        else:
            blocks.append(('para', line.strip()))
    return blocks


def strip_md_inline(text):
    """Strip bold/italic markers for plain text — docx handles styling separately."""
    import re
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def md_to_docx(input_md, output_docx):
    with open(input_md, 'r') as f:
        md_text = f.read()

    blocks = parse_md(md_text)
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = 914400  # 1 inch
        section.left_margin = section.right_margin = 914400

    skip_next_blank = False

    for i, (btype, content) in enumerate(blocks):
        if btype == 'h1':
            p = doc.add_heading(content, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.color.rgb = ACCENT
            p.runs[0].font.size = Pt(22)

        elif btype == 'h2':
            p = doc.add_heading(content.upper(), level=2)
            p.runs[0].font.color.rgb = ACCENT
            p.runs[0].font.size = Pt(11)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            add_hr(doc)

        elif btype == 'h3':
            # Job title line
            p = doc.add_paragraph()
            r = p.add_run(strip_md_inline(content))
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = ACCENT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)

        elif btype == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(strip_md_inline(content)).font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(2)

        elif btype == 'hr':
            add_hr(doc)

        elif btype == 'para' and content:
            # Check for bold-prefixed skill rows (e.g. **Label:** content)
            import re
            m = re.match(r'\*\*(.+?)\*\*[:：]\s*(.*)', content)
            if m:
                p = doc.add_paragraph()
                r1 = p.add_run(m.group(1) + ': ')
                r1.bold = True
                r1.font.size = Pt(9.5)
                rest = m.group(2).replace(' | ', '  |  ')
                r2 = p.add_run(strip_md_inline(rest))
                r2.font.size = Pt(9.5)
                p.paragraph_format.space_after = Pt(3)
            else:
                p = doc.add_paragraph(strip_md_inline(content))
                p.runs[0].font.size = Pt(9.5)
                p.paragraph_format.space_after = Pt(3)

        elif btype == 'blank':
            pass  # skip blank lines — spacing handled via paragraph_format

    doc.save(output_docx)
    print(f"Saved: {output_docx}")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python3 cv_to_docx.py input.md output.docx")
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
